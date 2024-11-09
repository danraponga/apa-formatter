from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def create_test_document():
    doc = Document()

    # Set incorrect margins (1.5 inches instead of 1 inch)
    section = doc.sections[0]
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)

    # Add header with violations
    header = section.header
    # Lowercase running head (violation)
    header_para = header.paragraphs[0]
    header_para.text = "running head: test document"
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Missing page number (violation)

    # Add paragraphs with font violations
    para1 = doc.add_paragraph("This paragraph uses Arial font.")
    run1 = para1.runs[0]
    run1.font.name = "Arial"  # Wrong font
    run1.font.size = Pt(12)  # Correct size

    para2 = doc.add_paragraph("This paragraph uses correct font but wrong size.")
    run2 = para2.runs[0]
    run2.font.name = "Times New Roman"  # Correct font
    run2.font.size = Pt(14)  # Wrong size

    para3 = doc.add_paragraph(
        "This paragraph has correct font and size but wrong spacing."
    )
    run3 = para3.runs[0]
    run3.font.name = "Times New Roman"
    run3.font.size = Pt(12)
    para3.paragraph_format.line_spacing = 1.5  # Wrong spacing (should be 2.0)

    # Add a paragraph with mixed formatting violations
    para4 = doc.add_paragraph()
    run4_1 = para4.add_run("This paragraph has ")
    run4_1.font.name = "Times New Roman"
    run4_1.font.size = Pt(12)

    run4_2 = para4.add_run("multiple formatting ")
    run4_2.font.name = "Arial"  # Wrong font
    run4_2.font.size = Pt(14)  # Wrong size

    run4_3 = para4.add_run("violations in different runs.")
    run4_3.font.name = "Calibri"  # Wrong font
    run4_3.font.size = Pt(11)  # Wrong size

    # Add a paragraph with correct formatting for comparison
    para5 = doc.add_paragraph("This paragraph follows all formatting rules correctly.")
    run5 = para5.runs[0]
    run5.font.name = "Times New Roman"
    run5.font.size = Pt(12)
    para5.paragraph_format.line_spacing = 2.0

    return doc
